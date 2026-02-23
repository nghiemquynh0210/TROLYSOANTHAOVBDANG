"""
Script để trích xuất nội dung từ các file DOCX trong thư mục HUONGDAN MAUCHI BO
"""
import os
import json
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Cần cài đặt python-docx: pip install python-docx")
    exit(1)

def extract_text_from_docx(file_path):
    """Trích xuất text từ file DOCX"""
    try:
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Trích xuất text từ tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Lỗi khi đọc {file_path}: {e}")
        return None

def analyze_directory(base_path):
    """Phân tích toàn bộ thư mục và trích xuất nội dung"""
    base_path = Path(base_path)
    results = {}
    
    # Các quy trình chính
    processes = [
        "GIAM SAT HOAN CHINH",
        "KIEM TRA HOAN CHINH", 
        "QUY TRINH GIAI QUYET TO CAO",
        "QUY TRINH KT DAU HIEU VI PHAM",
        "QUY TRINH THI HANH KY LUAT"
    ]
    
    for process in processes:
        process_path = base_path / process
        if not process_path.exists():
            print(f"Không tìm thấy thư mục: {process}")
            continue
        
        print(f"\n{'='*60}")
        print(f"Đang xử lý: {process}")
        print(f"{'='*60}")
        
        process_data = {
            "name": process,
            "files": []
        }
        
        # Lấy tất cả file DOCX và DOC
        doc_files = list(process_path.glob("*.docx")) + list(process_path.glob("*.doc"))
        doc_files = [f for f in doc_files if not f.name.startswith("~$")]  # Bỏ qua temp files
        doc_files.sort()
        
        for doc_file in doc_files:
            print(f"  Đang đọc: {doc_file.name}")
            content = extract_text_from_docx(doc_file)
            
            if content:
                file_data = {
                    "filename": doc_file.name,
                    "content": content,
                    "length": len(content),
                    "lines": content.count('\n') + 1
                }
                process_data["files"].append(file_data)
                print(f"    ✓ Đã trích xuất {len(content)} ký tự, {file_data['lines']} dòng")
            else:
                print(f"    ✗ Không thể đọc file")
        
        results[process] = process_data
        print(f"\nTổng số file đã xử lý: {len(process_data['files'])}")
    
    return results

def save_results(results, output_dir):
    """Lưu kết quả vào file"""
    output_dir = Path(output_dir)
    output_dir.mkdir(exist_ok=True)
    
    # Lưu toàn bộ vào JSON
    json_path = output_dir / "all_documents.json"
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print(f"\n✓ Đã lưu JSON: {json_path}")
    
    # Lưu từng quy trình vào file riêng
    for process_name, process_data in results.items():
        # Tạo file markdown cho từng quy trình
        md_path = output_dir / f"{process_name.replace(' ', '_')}.md"
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(f"# {process_name}\n\n")
            f.write(f"Tổng số file: {len(process_data['files'])}\n\n")
            f.write("---\n\n")
            
            for file_data in process_data['files']:
                f.write(f"## {file_data['filename']}\n\n")
                f.write(f"**Thông tin:**\n")
                f.write(f"- Độ dài: {file_data['length']} ký tự\n")
                f.write(f"- Số dòng: {file_data['lines']}\n\n")
                f.write("**Nội dung:**\n\n")
                f.write("```\n")
                f.write(file_data['content'])
                f.write("\n```\n\n")
                f.write("---\n\n")
        
        print(f"✓ Đã lưu Markdown: {md_path}")

def main():
    # Đường dẫn thư mục gốc
    base_path = r"C:\Users\Admin\Downloads\HUONGDAN MAUCHI BO"
    output_dir = r"C:\Users\Admin\Downloads\TROLYSOANTHAOVBDANG\extracted_templates"
    
    print("="*60)
    print("TRÍCH XUẤT NỘI DUNG TỪ CÁC FILE MẪU VĂN BẢN ĐẢNG")
    print("="*60)
    
    # Phân tích và trích xuất
    results = analyze_directory(base_path)
    
    # Lưu kết quả
    save_results(results, output_dir)
    
    # Thống kê tổng quan
    print("\n" + "="*60)
    print("THỐNG KÊ TỔNG QUAN")
    print("="*60)
    
    total_files = 0
    total_chars = 0
    
    for process_name, process_data in results.items():
        num_files = len(process_data['files'])
        total_files += num_files
        
        chars = sum(f['length'] for f in process_data['files'])
        total_chars += chars
        
        print(f"\n{process_name}:")
        print(f"  - Số file: {num_files}")
        print(f"  - Tổng ký tự: {chars:,}")
    
    print(f"\n{'='*60}")
    print(f"TỔNG CỘNG:")
    print(f"  - Tổng số file: {total_files}")
    print(f"  - Tổng số ký tự: {total_chars:,}")
    print(f"{'='*60}")

if __name__ == "__main__":
    main()
